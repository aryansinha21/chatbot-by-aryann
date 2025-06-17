import streamlit as st
import requests
import fitz  # PyMuPDF
import docx

# Streamlit page config
st.set_page_config(
    page_title="AI File Chatbot by Aryann Sinha",
    page_icon="📚",
    layout="centered",
    initial_sidebar_state="collapsed"
)

# Header Section
st.markdown("""
    <style>
    .title {
        text-align: center;
        font-size: 2.5em;
        color: #4CAF50;
        font-weight: bold;
    }
    .subtitle {
        text-align: center;
        font-size: 1.2em;
        margin-bottom: 30px;
    }
    </style>
    <div class="title">📚 Chat With Your File</div>
    <div class="subtitle">Upload a document and ask questions about it using AI.<br><b>Made by Aryann Sinha</b></div>
""", unsafe_allow_html=True)

# Upload section
file = st.file_uploader("📤 Upload PDF, DOCX, or TXT file", type=["pdf", "docx", "txt"])
question = st.text_input("🧠 Ask something about the file")

# Text extractor
def extract_text(file):
    if file.name.endswith(".pdf"):
        doc = fitz.open(stream=file.read(), filetype="pdf")
        return "\n".join([page.get_text() for page in doc])
    elif file.name.endswith(".docx"):
        doc = docx.Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    elif file.name.endswith(".txt"):
        return file.read().decode("utf-8")
    return "Unsupported file format."

# AI querying using OpenRouter API
def query_ai(prompt):
    headers = {
        "Authorization": f"Bearer {st.secrets['api_key']}",
        "Content-Type": "application/json"
    }
    data = {
        "model": "mistralai/mistral-7b-instruct",  # or "meta-llama/llama-3-8b-instruct"
        "messages": [
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ]
    }
    response = requests.post("https://openrouter.ai/api/v1/chat/completions", headers=headers, json=data)
    return response.json()["choices"][0]["message"]["content"]

# Main Logic
if file and question:
    with st.spinner("Analyzing file and generating answer..."):
        content = extract_text(file)
        final_prompt = f"Answer based on the document content:\n\n{content}\n\nQ: {question}"
        try:
            answer = query_ai(final_prompt)
            st.success(answer)
        except Exception as e:
            st.error(f"Error: {str(e)}")

# Footer
st.markdown("""
    <hr style="margin-top: 50px;">
    <p style="text-align: center; font-size: 0.9em; color: gray;">© 2025 Aryann Sinha. All rights reserved.</p>
""", unsafe_allow_html=True)
import streamlit as st
import requests
import fitz  # PyMuPDF
import docx

# Streamlit page config
st.set_page_config(
    page_title="AI File Chatbot by Aryann Sinha",
    page_icon="📚",
    layout="centered",
    initial_sidebar_state="collapsed"
)

# Header Section
st.markdown("""
    <style>
    .title {
        text-align: center;
        font-size: 2.5em;
        color: #4CAF50;
        font-weight: bold;
    }
    .subtitle {
        text-align: center;
        font-size: 1.2em;
        margin-bottom: 30px;
    }
    </style>
    <div class="title">📚 Chat With Your File</div>
    <div class="subtitle">Upload a document and ask questions about it using AI.<br><b>Made by Aryann Sinha</b></div>
""", unsafe_allow_html=True)

# Upload section
file = st.file_uploader("📤 Upload PDF, DOCX, or TXT file", type=["pdf", "docx", "txt"])
question = st.text_input("🧠 Ask something about the file")

# Text extractor
def extract_text(file):
    if file.name.endswith(".pdf"):
        doc = fitz.open(stream=file.read(), filetype="pdf")
        return "\n".join([page.get_text() for page in doc])
    elif file.name.endswith(".docx"):
        doc = docx.Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    elif file.name.endswith(".txt"):
        return file.read().decode("utf-8")
    return "Unsupported file format."

# AI querying using OpenRouter API
def query_ai(prompt):
    headers = {
        "Authorization": f"Bearer {st.secrets['api_key']}",
        "Content-Type": "application/json"
    }
    data = {
        "model": "mistralai/mistral-7b-instruct",  # or "meta-llama/llama-3-8b-instruct"
        "messages": [
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ]
    }
    response = requests.post("https://openrouter.ai/api/v1/chat/completions", headers=headers, json=data)
    return response.json()["choices"][0]["message"]["content"]

# Main Logic
if file and question:
    with st.spinner("Analyzing file and generating answer..."):
        content = extract_text(file)
        final_prompt = f"Answer based on the document content:\n\n{content}\n\nQ: {question}"
        try:
            answer = query_ai(final_prompt)
            st.success(answer)
        except Exception as e:
            st.error(f"Error: {str(e)}")

# Footer
st.markdown("""
    <hr style="margin-top: 50px;">
    <p style="text-align: center; font-size: 0.9em; color: gray;">© 2025 Aryann Sinha. All rights reserved.</p>
""", unsafe_allow_html=True)
